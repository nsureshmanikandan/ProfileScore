from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.schemas import ResumeRequest
from app.services.llm_service import generate_resume_text
from app.core.logging import get_logger
from app.core.telemetry import get_tracer
from opentelemetry.trace import SpanKind

logger = get_logger(__name__)
tracer = get_tracer("api.resume")

router = APIRouter(prefix="/api/resume", tags=["resume"])

# Brand colors
_BLUE   = RGBColor(0x1a, 0x56, 0xdb)   # section header blue
_DARK   = RGBColor(0x1e, 0x29, 0x3b)   # name / body dark
_GREY   = RGBColor(0x64, 0x74, 0x8b)   # contact / meta grey
_LINE   = RGBColor(0xca, 0xd5, 0xe2)   # divider colour


@router.post("/generate")
async def generate_resume(request: ResumeRequest):
    with tracer.start_as_current_span("api.generate_resume", kind=SpanKind.SERVER):
        logger.info("resume_generation_start", format=request.format)

        raw = request.analysis.raw_sections or {}
        # Prefer explicit fields; fall back to raw_sections (older analyses before schema update)
        name         = request.analysis.person_name         or raw.get("name", "")
        linkedin_url = request.analysis.person_linkedin_url or raw.get("linkedin_url", "")
        location     = request.analysis.person_location     or raw.get("location", "")

        analysis_dict = request.analysis.model_dump()
        resume_text = generate_resume_text(
            analysis_dict,
            request.target_role,
            person_name=name,
            linkedin_url=linkedin_url,
            location=location,
        )

        if request.format == "docx":
            docx_bytes = _build_docx(
                resume_text,
                name=name,
                linkedin_url=linkedin_url,
                location=location,
            )
            return StreamingResponse(
                io.BytesIO(docx_bytes),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=ProfileScore_Resume.docx"},
            )
        else:
            header = f"{request.analysis.person_name}\n"
            if request.analysis.person_linkedin_url:
                header += f"{request.analysis.person_linkedin_url}  "
            if request.analysis.person_location:
                header += f"{request.analysis.person_location}"
            header = header.strip() + "\n\n"
            return StreamingResponse(
                io.BytesIO((header + resume_text).encode("utf-8")),
                media_type="text/plain",
                headers={"Content-Disposition": "attachment; filename=ProfileScore_Resume.txt"},
            )


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _set_margins(doc: Document, top=Cm(1.5), bottom=Cm(1.5), left=Cm(1.8), right=Cm(1.8)):
    for section in doc.sections:
        section.top_margin = top
        section.bottom_margin = bottom
        section.left_margin = left
        section.right_margin = right


def _add_bottom_border(paragraph, color: str = "CAD5E2"):
    """Add a thin bottom border to a paragraph (used under section headings)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _para_spacing(paragraph, before: int = 0, after: int = 0, line_rule: str = "auto", line: int = 240):
    """Set paragraph spacing (twips: 240 = single, 276 = 1.15x)."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:lineRule'), line_rule)
    spacing.set(qn('w:line'), str(line))
    pPr.append(spacing)


def _add_name_header(doc: Document, name: str, linkedin_url: str, location: str):
    # Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_name, before=0, after=16)
    run = p_name.add_run(name if name else "Your Name")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _DARK
    run.font.name = "Calibri"

    # Contact line — row 1: email · phone · location
    contact_row1 = []
    contact_row1.append("your.email@gmail.com")
    contact_row1.append("+91-XXXXX-XXXXX")
    if location:
        contact_row1.append(location)

    p_c1 = doc.add_paragraph()
    p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_c1, before=0, after=6)
    r1 = p_c1.add_run("  ·  ".join(contact_row1))
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = _GREY
    r1.font.name = "Calibri"

    # Contact line — row 2: LinkedIn URL
    linkedin_display = linkedin_url if linkedin_url else "linkedin.com/in/your-profile"
    p_c2 = doc.add_paragraph()
    p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_c2, before=0, after=60)
    _add_bottom_border(p_c2, color="1A56DB")
    r2 = p_c2.add_run(linkedin_display)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = _BLUE
    r2.font.name = "Calibri"


def _add_section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    _para_spacing(p, before=160, after=40)
    _add_bottom_border(p)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = _BLUE
    run.font.name = "Calibri"


def _add_body_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
                   before: int = 0, after: int = 40):
    p = doc.add_paragraph()
    _para_spacing(p, before=before, after=after, line=264)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = _DARK
    run.font.name = "Calibri"
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    _para_spacing(p, before=0, after=30, line=252)
    run = p.add_run(text.lstrip("•·-– "))
    run.font.size = Pt(10.5)
    run.font.color.rgb = _DARK
    run.font.name = "Calibri"


def _build_docx(resume_text: str, name: str = "", linkedin_url: str = "", location: str = "") -> bytes:
    doc = Document()
    _set_margins(doc)

    # Default normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Header block ──
    _add_name_header(doc, name, linkedin_url, location)

    # ── Parse and render body ──
    lines = resume_text.split('\n')
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        i += 1

        if not stripped:
            continue

        # Section heading (ALL CAPS line with no leading symbols, not too long)
        if (stripped.isupper() and 3 < len(stripped) < 50
                and not stripped.startswith('•')
                and not re.match(r'^[A-Z\s]+\|', stripped)):
            _add_section_heading(doc, stripped)
            continue

        # Bullet point
        if stripped.startswith(('•', '-', '*', '·')):
            _add_bullet(doc, stripped)
            continue

        # Role/company line — "Company | Title | Year – Year"
        if '|' in stripped and re.search(r'\d{4}', stripped):
            _add_body_para(doc, stripped, bold=True, before=80, after=20)
            continue

        # Skills pipe-separated line
        if stripped.count('|') >= 3:
            p = _add_body_para(doc, stripped, before=0, after=40)
            continue

        # Regular body paragraph
        _add_body_para(doc, stripped, before=0, after=30)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
